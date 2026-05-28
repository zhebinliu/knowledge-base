"""会议纪要版面模板 API。

提供模板 CRUD、文件上传智能解析、模板渲染与导出。
路由注册在 main.py：prefix="/api/markup-templates"
"""
from __future__ import annotations

import io
import json
import logging
from datetime import datetime
from typing import Optional

from fastapi import APIRouter, Depends, File, Form, HTTPException, Query, UploadFile
from fastapi.responses import Response
from pydantic import BaseModel, Field
from sqlalchemy import select, delete as sql_delete
from sqlalchemy.ext.asyncio import AsyncSession

from models import get_session
from models.markup_template import MarkupTemplate
from models.meeting import Meeting, Requirement
from models.user import User
from services.auth import get_current_user

logger = logging.getLogger(__name__)

router = APIRouter()

_MAX_UPLOAD_MB = 20
_MAX_UPLOAD_BYTES = _MAX_UPLOAD_MB * 1024 * 1024

ALLOWED_EXTS = {".md", ".markdown", ".docx", ".doc", ".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"}


# ── Pydantic Schemas ──────────────────────────────────────────────────────

class TemplateCreate(BaseModel):
    name: str = Field(default="自定义模板", max_length=256)
    description: str = Field(default="")
    content: str = Field(default="", description="Markdown 内容")


class TemplateUpdate(BaseModel):
    name: Optional[str] = Field(default=None, max_length=256)
    description: Optional[str] = None
    content: Optional[str] = None


class TemplateRenderRequest(BaseModel):
    meeting_id: int


# ── 权限辅助 ──────────────────────────────────────────────────────────────

async def _load_template(template_id: int, db: AsyncSession) -> MarkupTemplate:
    tpl = await db.get(MarkupTemplate, template_id)
    if not tpl:
        raise HTTPException(status_code=404, detail="模板不存在")
    return tpl


# ── 文件解析辅助 ──────────────────────────────────────────────────────────

def _parse_markdown_file(content: bytes) -> str:
    """解析 .md/.markdown 文件为纯文本。"""
    for enc in ("utf-8", "utf-8-sig", "gbk", "gb2312", "latin-1"):
        try:
            return content.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return content.decode("utf-8", errors="replace")


def _parse_docx_file(content: bytes) -> str:
    """解析 .docx 文件为 Markdown。"""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(content))

        lines: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                lines.append("")
                continue
            # 根据样式推断标题级别
            style_name = (para.style.name if para.style else "").lower()
            if "heading 1" in style_name or style_name == "title":
                lines.append(f"# {text}")
            elif "heading 2" in style_name:
                lines.append(f"## {text}")
            elif "heading 3" in style_name:
                lines.append(f"### {text}")
            elif para.runs and para.runs[0].bold:
                lines.append(f"**{text}**")
            else:
                lines.append(text)

        # 提取表格
        for table in doc.tables:
            lines.append("")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
            lines.append("")

        return "\n".join(lines).strip() or "# 会议纪要"

    except ImportError:
        raise HTTPException(500, "python-docx 未安装，无法解析 Word 文档")
    except Exception as e:
        logger.exception("docx_parse_failed", error=str(e)[:200])
        raise HTTPException(400, f"Word 文档解析失败：{e}")


async def _parse_image_file(content: bytes, filename: str) -> str:
    """解析图片文件为 Markdown（通过多模态 LLM OCR）。"""
    import base64
    import mimetypes

    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "png"
    mime = mimetypes.guess_type(f"file.{ext}")[0] or f"image/{ext}"
    b64 = base64.b64encode(content).decode("ascii")
    data_url = f"data:{mime};base64,{b64}"

    prompt = (
        "你是一个专业的文档识别助手。请将这张图片中的文档内容提取出来，"
        "转换为结构清晰的 Markdown 格式。\n"
        "规则：\n"
        "1. 识别标题层级（# ## ###），保留原文结构\n"
        "2. 识别表格并转为 Markdown 表格\n"
        "3. 识别列表（有序/无序）\n"
        "4. 保留粗体、斜体等格式标记\n"
        "5. 忽略水印和无关背景\n"
        "6. 如果有占位符标记如 {会议名称} [日期] 等，保留原样\n"
        "7. 只输出 Markdown 内容，不要加任何解释"
    )

    try:
        from services.model_router import model_router
        messages = [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": data_url}},
                ],
            },
        ]
        result, model = await model_router.chat_with_routing(
            task="pdf_ocr",  # 使用 mimo-v2-omni 多模态模型做图片 OCR
            messages=messages,
            temperature=0.1,
            max_tokens=8000,
        )
        logger.info("image_ocr_done", filename=filename, model=model, length=len(result))
        return result.strip()
    except Exception as e:
        logger.exception("image_ocr_failed", filename=filename, error=str(e)[:200])
        raise HTTPException(500, f"图片识别失败：{e}")


async def _parse_uploaded_file(content: bytes, filename: str) -> tuple[str, str]:
    """解析上传文件，返回 (markdown_content, source_format)。"""
    ext = (filename.rsplit(".", 1)[-1].lower() if "." in filename else "").lower()
    if not ext:
        raise HTTPException(400, f"无法识别文件格式：{filename}")

    if ext in ("md", "markdown", "txt"):
        return _parse_markdown_file(content), "markdown"
    elif ext in ("docx", "doc"):
        if ext == "doc":
            raise HTTPException(400, "暂不支持 .doc 格式，请转换为 .docx 后上传")
        return _parse_docx_file(content), "docx"
    elif ext in ("png", "jpg", "jpeg", "gif", "webp", "bmp"):
        result = await _parse_image_file(content, filename)
        return result, "image"
    else:
        raise HTTPException(400, f"不支持的文件格式：.{ext}。支持：md, markdown, docx, png, jpg, jpeg, gif, webp")


# ── 模板渲染 ─────────────────────────────────────────────────────────────

# 可用占位符说明（前端展示用）
PLACEHOLDER_HELP = {
    "{{title}}": "会议标题",
    "{{date}}": "会议日期",
    "{{time}}": "会议时间",
    "{{location}}": "会议地点",
    "{{host}}": "主持人",
    "{{recorder}}": "记录人",
    "{{attendees}}": "参会人员列表",
    "{{summary}}": "会议摘要",
    "{{key_points}}": "关键议题（Markdown 列表）",
    "{{decisions}}": "决议事项（Markdown 列表）",
    "{{action_items}}": "待办事项（Markdown 表格）",
    "{{unresolved}}": "未决问题（Markdown 列表）",
    "{{requirements}}": "需求清单（Markdown 表格）",
    "{{stakeholders}}": "干系人列表",
    "{{transcript_summary}}": "转录摘要（前 2000 字）",
}


def _format_key_points_md(items: list) -> str:
    if not items:
        return "_(暂无关键议题)_"
    lines = []
    for i, kp in enumerate(items, start=1):
        if isinstance(kp, dict):
            topic = kp.get("topic", "")
            content = kp.get("content", "")
            lines.append(f"{i}. **{topic}**：{content}")
        else:
            lines.append(f"{i}. {kp}")
    return "\n".join(lines)


def _format_decisions_md(items: list) -> str:
    if not items:
        return "_(暂无决议)_"
    lines = []
    for d in items:
        if isinstance(d, dict):
            content = d.get("content", "")
            owner = d.get("owner", "")
            line = f"- {content}"
            if owner:
                line += f"（负责人：{owner}）"
            lines.append(line)
        else:
            lines.append(f"- {d}")
    return "\n".join(lines)


def _format_action_items_md(items: list) -> str:
    if not items:
        return "_(暂无待办)_"
    lines = ["| 序号 | 事项 | 负责人 | 截止日期 | 备注 |", "| --- | --- | --- | --- | --- |"]
    for i, a in enumerate(items, start=1):
        if isinstance(a, dict):
            task = a.get("task", "")
            owner = a.get("owner", "")
            deadline = a.get("deadline", "")
            remark = a.get("remark", "") or a.get("priority", "")
            lines.append(f"| {i} | {task} | {owner} | {deadline} | {remark} |")
        else:
            lines.append(f"| {i} | {a} | | | |")
    return "\n".join(lines)


def _format_unresolved_md(items: list) -> str:
    if not items:
        return "_(暂无)_"
    lines = []
    for u in items:
        if isinstance(u, dict):
            issue = u.get("issue", "")
            reason = u.get("reason", "")
            line = f"- {issue}"
            if reason:
                line += f"（原因：{reason}）"
            lines.append(line)
        else:
            lines.append(f"- {u}")
    return "\n".join(lines)


def _format_requirements_md(reqs: list[dict]) -> str:
    if not reqs:
        return "_(暂无需求)_"
    lines = ["| 编号 | 模块 | 描述 | 优先级 | 来源 | 状态 |", "| --- | --- | --- | --- | --- | --- |"]
    for r in reqs:
        lines.append(
            f"| {r.get('req_id', '')} | {r.get('module', '')} | {r.get('description', '')} "
            f"| {r.get('priority', '')} | {r.get('speaker', '') or r.get('source', '')} "
            f"| {r.get('status', '')} |"
        )
    return "\n".join(lines)


def _format_attendees_md(attendees: list) -> str:
    if not attendees:
        return ""
    return "、".join(str(a) for a in attendees)


def _format_stakeholders_md(smap: dict | None) -> str:
    if not smap or not isinstance(smap, dict):
        return "_(暂无干系人信息)_"
    holders = smap.get("stakeholders") or []
    if not holders:
        return "_(暂无干系人信息)_"
    lines = ["| 姓名 | 角色 | 组织 | 立场 |", "| --- | --- | --- | --- |"]
    for h in holders:
        if not isinstance(h, dict):
            continue
        lines.append(
            f"| {h.get('name', '')} | {h.get('role', '')} "
            f"| {h.get('organization', '')} | {h.get('side', '')} |"
        )
    return "\n".join(lines)


def render_template(template_content: str, meeting_data: dict) -> str:
    """将模板中的占位符替换为会议数据，返回渲染后的 Markdown。"""
    mm = meeting_data.get("meeting_minutes") or {}
    if not isinstance(mm, dict):
        mm = {}

    result = template_content

    # 基础字段
    result = result.replace("{{title}}", str(meeting_data.get("title") or ""))
    result = result.replace("{{summary}}", str(mm.get("summary") or ""))

    # 日期时间
    start_time = meeting_data.get("start_time")
    if start_time:
        if isinstance(start_time, str):
            dt = start_time[:19].replace("T", " ")
        else:
            dt = str(start_time)[:19].replace("T", " ")
        result = result.replace("{{date}}", dt[:10])
        result = result.replace("{{time}}", dt[11:16] if len(dt) > 11 else "")
    else:
        result = result.replace("{{date}}", "")
        result = result.replace("{{time}}", "")

    # 会议元信息
    result = result.replace("{{location}}", str(mm.get("meeting_location") or ""))
    result = result.replace("{{host}}", str(mm.get("meeting_host") or ""))
    result = result.replace("{{recorder}}", str(mm.get("meeting_recorder") or ""))

    # 参会人员
    attendees = mm.get("attendees") or []
    result = result.replace("{{attendees}}", _format_attendees_md(attendees))

    # 结构化内容
    result = result.replace("{{key_points}}", _format_key_points_md(mm.get("key_points") or []))
    result = result.replace("{{decisions}}", _format_decisions_md(mm.get("decisions") or []))
    result = result.replace("{{action_items}}", _format_action_items_md(mm.get("action_items") or []))
    result = result.replace("{{unresolved}}", _format_unresolved_md(mm.get("unresolved") or []))

    # 需求清单
    reqs = meeting_data.get("requirements") or []
    result = result.replace("{{requirements}}", _format_requirements_md(reqs))

    # 干系人
    smap = meeting_data.get("stakeholder_map")
    result = result.replace("{{stakeholders}}", _format_stakeholders_md(smap))

    # 转录摘要
    polished = meeting_data.get("polished_transcript") or meeting_data.get("raw_transcript") or ""
    transcript_summary = polished[:2000] if polished else "_(暂无转录内容)_"
    result = result.replace("{{transcript_summary}}", transcript_summary)

    return result


# ── 导出：渲染后 Markdown → DOCX ──────────────────────────────────────────

def _render_markdown_to_docx(md_content: str) -> bytes:
    """将 Markdown 文本渲染为 DOCX bytes。"""
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Inches, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
    except ImportError:
        raise HTTPException(500, "python-docx 未安装")

    doc = DocxDocument()

    # 设置默认样式
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    # 配置标题样式
    for level, (size, bold) in {1: (18, True), 2: (15, True), 3: (12.5, True)}.items():
        try:
            h_style = doc.styles[f"Heading {level}"]
            h_style.font.size = Pt(size)
            h_style.font.bold = bold
            h_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        except KeyError:
            pass

    lines = md_content.split("\n")
    i = 0
    in_table = False
    table_data: list[list[str]] = []
    table_header: list[str] = []

    while i < len(lines):
        line = lines[i]

        # 空行
        if not line.strip():
            if in_table:
                # 表格结束
                _build_table(doc, table_data, table_header)
                in_table = False
                table_data = []
                table_header = []
            doc.add_paragraph("")
            i += 1
            continue

        # 标题
        if line.startswith("# "):
            if in_table:
                _build_table(doc, table_data, table_header)
                in_table = False
                table_data = []
                table_header = []
            h = doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        elif line.startswith("## "):
            if in_table:
                _build_table(doc, table_data, table_header)
                in_table = False
                table_data = []
                table_header = []
            h = doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        elif line.startswith("### "):
            if in_table:
                _build_table(doc, table_data, table_header)
                in_table = False
                table_data = []
                table_header = []
            h = doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        # 表格
        if line.strip().startswith("|") and line.strip().endswith("|"):
            # 跳过分隔行
            if all(c in "|-: " for c in line.strip()):
                i += 1
                continue
            if not in_table:
                in_table = True
                table_header = [c.strip() for c in line.split("|")[1:-1]]
            else:
                cells = [c.strip() for c in line.split("|")[1:-1]]
                table_data.append(cells)
            i += 1
            continue

        # 列表项
        if line.strip().startswith("- ") or line.strip().startswith("* "):
            text = line.strip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_runs(p, text)
            i += 1
            continue

        # 有序列表
        stripped = line.strip()
        if stripped and stripped[0].isdigit() and ". " in stripped[:6]:
            dot_pos = stripped.index(". ")
            text = stripped[dot_pos + 2:]
            p = doc.add_paragraph(style="List Number")
            _add_formatted_runs(p, text)
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        _add_formatted_runs(p, line)
        i += 1

    # 处理最后的表格
    if in_table:
        _build_table(doc, table_data, table_header)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_formatted_runs(paragraph, text: str):
    """添加带粗体/斜体格式的 run 到段落。"""
    import re

    # 简单处理：**粗体** 和 *斜体*
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def _build_table(doc, data: list[list[str]], header: list[str]):
    """在文档中创建表格。"""
    if not header and not data:
        return
    all_rows = ([header] if header else []) + data
    if not all_rows:
        return
    ncols = max(len(r) for r in all_rows)
    table = doc.add_table(rows=len(all_rows), cols=ncols)
    table.style = "Light Grid Accent 1"

    for ri, row_data in enumerate(all_rows):
        for ci in range(ncols):
            cell = table.cell(ri, ci)
            cell.text = row_data[ci] if ci < len(row_data) else ""
            # 表头加粗
            if ri == 0 and header:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


# ── Endpoints ────────────────────────────────────────────────────────────

@router.get("", response_model=list[dict])
async def list_templates(
    db: AsyncSession = Depends(get_session),
):
    """列出所有版面模板（按更新时间倒序）。"""
    result = await db.execute(
        select(MarkupTemplate).order_by(MarkupTemplate.updated_at.desc())
    )
    templates = list(result.scalars().all())
    return [t.to_dict() for t in templates]


@router.get("/placeholders", response_model=dict)
async def get_placeholders():
    """返回可用占位符说明，供前端展示。"""
    return PLACEHOLDER_HELP


@router.get("/{template_id}", response_model=dict)
async def get_template(
    template_id: int,
    db: AsyncSession = Depends(get_session),
):
    """按 ID 获取单个模板。"""
    tpl = await _load_template(template_id, db)
    return tpl.to_dict()


@router.post("", response_model=dict, status_code=201)
async def create_template(
    payload: TemplateCreate,
    db: AsyncSession = Depends(get_session),
    user: User = Depends(get_current_user),
):
    """手动创建模板（输入 Markdown）。"""
    tpl = MarkupTemplate(
        name=payload.name,
        description=payload.description,
        content=payload.content,
        category="user_upload",
        source_format="markdown",
        is_builtin=False,
    )
    db.add(tpl)
    await db.commit()
    await db.refresh(tpl)
    logger.info("markup_template_created", id=tpl.id, name=tpl.name, user=user.username)
    return tpl.to_dict()


@router.post("/upload", response_model=dict, status_code=201)
async def upload_template(
    file: UploadFile = File(...),
    name: Optional[str] = Form(None),
    description: Optional[str] = Form(""),
    db: AsyncSession = Depends(get_session),
    user: User = Depends(get_current_user),
):
    """上传模板文件（.md / .docx / 图片），自动解析为 Markdown 并保存。"""
    filename = (file.filename or "template").strip()

    # 校验扩展名
    ext = ("." + filename.rsplit(".", 1)[-1].lower()) if "." in filename else ""
    if ext not in ALLOWED_EXTS:
        raise HTTPException(
            400,
            f"不支持的文件格式：{ext}。支持：{', '.join(sorted(ALLOWED_EXTS))}",
        )

    # 读文件内容 + 大小校验
    content = await file.read()
    if not content:
        raise HTTPException(400, "上传文件为空")
    if len(content) > _MAX_UPLOAD_BYTES:
        raise HTTPException(413, f"文件超过 {_MAX_UPLOAD_MB} MB 限制")

    # 解析
    md_content, source = await _parse_uploaded_file(content, filename)

    # 保存
    tpl = MarkupTemplate(
        name=name or filename.rsplit(".", 1)[0],
        description=description or "",
        content=md_content,
        category="user_upload",
        source_format=source,
        is_builtin=False,
    )
    db.add(tpl)
    await db.commit()
    await db.refresh(tpl)
    logger.info(
        "markup_template_uploaded",
        id=tpl.id,
        name=tpl.name,
        source=source,
        user=user.username,
        bytes=len(md_content),
    )
    return tpl.to_dict()


@router.patch("/{template_id}", response_model=dict)
async def update_template(
    template_id: int,
    payload: TemplateUpdate,
    db: AsyncSession = Depends(get_session),
    user: User = Depends(get_current_user),
):
    """更新模板名称、描述或内容。"""
    tpl = await _load_template(template_id, db)
    if payload.name is not None:
        tpl.name = payload.name
    if payload.description is not None:
        tpl.description = payload.description
    if payload.content is not None:
        tpl.content = payload.content
    await db.commit()
    await db.refresh(tpl)
    logger.info("markup_template_updated", id=tpl.id, user=user.username)
    return tpl.to_dict()


@router.delete("/{template_id}", status_code=204)
async def delete_template(
    template_id: int,
    db: AsyncSession = Depends(get_session),
    user: User = Depends(get_current_user),
):
    """删除模板（内置模板不可删除）。"""
    tpl = await _load_template(template_id, db)
    if tpl.is_builtin:
        raise HTTPException(400, "内置模板不可删除")
    await db.delete(tpl)
    await db.commit()
    logger.info("markup_template_deleted", id=template_id, name=tpl.name, user=user.username)
    return None


@router.post("/{template_id}/render", response_model=dict)
async def render_template_with_meeting(
    template_id: int,
    body: TemplateRenderRequest,
    db: AsyncSession = Depends(get_session),
    user: User = Depends(get_current_user),
):
    """用指定模板渲染某场会议的数据，返回渲染后 Markdown。

    前端可预览后决定是否导出。
    """
    tpl = await _load_template(template_id, db)

    # 加载会议数据
    m = await db.get(Meeting, body.meeting_id)
    if not m:
        raise HTTPException(404, "会议不存在")

    # 查需求
    reqs = (await db.scalars(
        select(Requirement).where(Requirement.meeting_id == m.id).order_by(Requirement.id)
    )).all()

    meeting_data = {
        "title": m.title,
        "start_time": m.start_time.isoformat() if m.start_time else None,
        "raw_transcript": m.raw_transcript or "",
        "polished_transcript": m.polished_transcript or "",
        "meeting_minutes": m.meeting_minutes,
        "stakeholder_map": m.stakeholder_map,
        "requirements": [
            {
                "req_id": r.req_id,
                "module": r.module,
                "description": r.description,
                "priority": r.priority,
                "source": r.source,
                "speaker": r.speaker,
                "status": r.status,
            }
            for r in reqs
        ],
    }

    rendered = render_template(tpl.content, meeting_data)

    logger.info(
        "markup_template_rendered",
        template_id=template_id,
        meeting_id=body.meeting_id,
        user=user.username,
        length=len(rendered),
    )
    return {
        "template_id": template_id,
        "template_name": tpl.name,
        "meeting_id": body.meeting_id,
        "meeting_title": m.title,
        "rendered": rendered,
    }


@router.post("/{template_id}/export-docx")
async def export_template_docx(
    template_id: int,
    body: TemplateRenderRequest,
    db: AsyncSession = Depends(get_session),
    user: User = Depends(get_current_user),
):
    """用指定模板渲染会议数据并导出为 DOCX 文件。"""
    from urllib.parse import quote

    tpl = await _load_template(template_id, db)
    m = await db.get(Meeting, body.meeting_id)
    if not m:
        raise HTTPException(404, "会议不存在")

    reqs = (await db.scalars(
        select(Requirement).where(Requirement.meeting_id == m.id).order_by(Requirement.id)
    )).all()

    meeting_data = {
        "title": m.title,
        "start_time": m.start_time.isoformat() if m.start_time else None,
        "raw_transcript": m.raw_transcript or "",
        "polished_transcript": m.polished_transcript or "",
        "meeting_minutes": m.meeting_minutes,
        "stakeholder_map": m.stakeholder_map,
        "requirements": [
            {
                "req_id": r.req_id,
                "module": r.module,
                "description": r.description,
                "priority": r.priority,
                "source": r.source,
                "speaker": r.speaker,
                "status": r.status,
            }
            for r in reqs
        ],
    }

    rendered_md = render_template(tpl.content, meeting_data)

    try:
        docx_bytes = _render_markdown_to_docx(rendered_md)
    except Exception as e:
        logger.exception("markup_template_docx_failed", template_id=template_id, error=str(e)[:200])
        raise HTTPException(500, f"生成 DOCX 失败：{e}")

    safe_name = quote(f"{m.title or '会议纪要'}-{tpl.name}.docx")
    ascii_name = "".join(c for c in (m.title or "meeting_minutes") if ord(c) < 128 and c not in '\\/:*?"<>|')
    ascii_name = (ascii_name.strip()[:50] or "meeting_minutes") + ".docx"

    logger.info(
        "markup_template_exported_docx",
        template_id=template_id,
        meeting_id=body.meeting_id,
        user=user.username,
        bytes=len(docx_bytes),
    )
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": (
                f"attachment; filename={quote(ascii_name)}; "
                f"filename*=UTF-8''{safe_name}"
            ),
        },
    )


@router.post("/{template_id}/export-md")
async def export_template_md(
    template_id: int,
    body: TemplateRenderRequest,
    db: AsyncSession = Depends(get_session),
    user: User = Depends(get_current_user),
):
    """用指定模板渲染会议数据并导出为 Markdown 文件。"""
    from urllib.parse import quote

    tpl = await _load_template(template_id, db)
    m = await db.get(Meeting, body.meeting_id)
    if not m:
        raise HTTPException(404, "会议不存在")

    reqs = (await db.scalars(
        select(Requirement).where(Requirement.meeting_id == m.id).order_by(Requirement.id)
    )).all()

    meeting_data = {
        "title": m.title,
        "start_time": m.start_time.isoformat() if m.start_time else None,
        "raw_transcript": m.raw_transcript or "",
        "polished_transcript": m.polished_transcript or "",
        "meeting_minutes": m.meeting_minutes,
        "stakeholder_map": m.stakeholder_map,
        "requirements": [
            {
                "req_id": r.req_id,
                "module": r.module,
                "description": r.description,
                "priority": r.priority,
                "source": r.source,
                "speaker": r.speaker,
                "status": r.status,
            }
            for r in reqs
        ],
    }

    rendered_md = render_template(tpl.content, meeting_data)

    safe_name = quote(f"{m.title or '会议纪要'}-{tpl.name}.md")
    ascii_name = "".join(c for c in (m.title or "meeting_minutes") if ord(c) < 128 and c not in '\\/:*?"<>|')
    ascii_name = (ascii_name.strip()[:50] or "meeting_minutes") + ".md"

    return Response(
        content=rendered_md.encode("utf-8"),
        media_type="text/markdown; charset=utf-8",
        headers={
            "Content-Disposition": (
                f"attachment; filename={quote(ascii_name)}; "
                f"filename*=UTF-8''{safe_name}"
            ),
        },
    )
