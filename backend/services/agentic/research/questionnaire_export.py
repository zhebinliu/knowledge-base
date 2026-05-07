"""会前调研问卷按角色导出 — Word / Excel / 打印 HTML 三种格式。

设计:
- 仅导出 phase='pre_meeting' 的题目(会中题由 PM 现场把控,不发给客户)
- 顾问指定角色(executive / dept_head / frontline / it / all)
- 输出**纯空白模板**:仅有题干 + 选项,客户拿到从零填(不预填已答)
- 题目按 ltc_module_key 分组(章节),组内保持原顺序
"""
from __future__ import annotations

import io
import re
from typing import Iterable

from .questionnaire_schema import AUDIENCE_ROLE_LABELS, VALID_AUDIENCE_ROLES


ROLE_ALL = "all"


def filter_items(
    items: list[dict],
    role: str,
) -> list[dict]:
    """筛会前 + 命中角色的题。role='all' 时只筛 phase。"""
    out: list[dict] = []
    for it in items:
        if (it.get("phase") or "in_meeting") != "pre_meeting":
            continue
        if role != ROLE_ALL:
            roles = it.get("audience_roles") or []
            if role not in roles:
                continue
        out.append(it)
    return out


def group_by_ltc_module(items: list[dict]) -> list[tuple[str, list[dict]]]:
    """按 ltc_module_key 分组,保留原顺序。"""
    groups: dict[str, list[dict]] = {}
    order: list[str] = []
    for it in items:
        k = it.get("ltc_module_key") or "_uncategorized"
        if k not in groups:
            groups[k] = []
            order.append(k)
        groups[k].append(it)
    return [(k, groups[k]) for k in order]


def role_label(role: str) -> str:
    return "全部角色" if role == ROLE_ALL else AUDIENCE_ROLE_LABELS.get(role, role)


def export_filename(project_name: str, role: str, ext: str) -> str:
    """规范化文件名:<项目>_会前调研问卷_<角色>.<ext>"""
    safe_proj = re.sub(r"[\\/:*?\"<>|]", "_", (project_name or "项目")).strip() or "项目"
    return f"{safe_proj}_会前调研问卷_{role_label(role)}.{ext}"


# ── Markdown / 纯文本格式器(给 docx + html + xlsx 共用)─────────────────────

def _option_lines(item: dict) -> list[str]:
    """根据题型生成「客户填空区」描述。"""
    t = item.get("type")
    options = item.get("options") or []
    if t == "single":
        out = []
        for o in options:
            label = o.get("label", "")
            out.append(f"☐ {label}")
        return out
    if t in ("multi", "node_pick"):
        out = []
        for o in options:
            label = o.get("label", "")
            out.append(f"☐ {label}")
        return out
    if t == "rating":
        scale = item.get("rating_scale", 5)
        return [f"评分 1 - {scale}(请圈选):" + "  ".join(str(i) for i in range(1, scale + 1))]
    if t == "number":
        unit = item.get("number_unit", "")
        suffix = f" ({unit})" if unit else ""
        return [f"请填数值{suffix}:____________________"]
    # text 类
    return [
        "答:_____________________________________________________________",
        "    _____________________________________________________________",
        "    _____________________________________________________________",
    ]


# ── DOCX ────────────────────────────────────────────────────────────────────

def export_docx(
    *,
    project_name: str,
    role: str,
    items: list[dict],
    ltc_label_lookup: dict[str, str] | None = None,
) -> bytes:
    """生成会前调研问卷 .docx 字节流。"""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 封面
    title = doc.add_heading(f"{project_name or '项目'} · 会前调研问卷", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"受访角色:{role_label(role)}").bold = True

    # 引导语
    intro = doc.add_paragraph()
    intro.add_run(
        "尊敬的同事您好,本问卷用于在调研会议前了解贵司的现状,帮助实施团队"
        "提前准备针对性的访谈话题。请您按当前真实情况勾选 / 填写;选项不全可在「其他」"
        "处补充,实在不适用的题可勾「不适用」。约 5-10 分钟可完成,会上我们会基于"
        "您的回答展开深入讨论。感谢配合!"
    )

    # 按 LTC 模块分组
    groups = group_by_ltc_module(items)
    if not groups:
        doc.add_paragraph().add_run("(本角色暂无会前题目)").italic = True
    else:
        global_idx = 0
        for ltc_key, sub_items in groups:
            module_label = (ltc_label_lookup or {}).get(ltc_key, ltc_key)
            doc.add_heading(module_label, level=1)
            for it in sub_items:
                global_idx += 1
                # 题干(N. 题目 + 必答 *)
                p = doc.add_paragraph()
                p.add_run(f"{global_idx}. ").bold = True
                p.add_run(it.get("question", ""))
                if it.get("required"):
                    star = p.add_run("  *")
                    star.bold = True
                # hint
                if it.get("hint"):
                    h = doc.add_paragraph()
                    h.add_run(it["hint"]).italic = True
                # 选项 / 填空
                for line in _option_lines(it):
                    doc.add_paragraph(line, style="List Bullet" if line.startswith("☐") else None)
                doc.add_paragraph()  # 留白

    # 末尾感谢
    doc.add_page_break()
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = end.add_run("感谢您的耐心填写,期待会上深入交流!")
    r.bold = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── XLSX ────────────────────────────────────────────────────────────────────

def export_xlsx(
    *,
    project_name: str,
    role: str,
    items: list[dict],
    ltc_label_lookup: dict[str, str] | None = None,
) -> bytes:
    """生成会前调研问卷 .xlsx 字节流。题目一行,选项展开到独立单元格,客户答案列留空。"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = Workbook()
    ws = wb.active
    ws.title = "会前调研问卷"

    # 标题行
    ws["A1"] = f"{project_name or '项目'} · 会前调研问卷 · {role_label(role)}"
    ws["A1"].font = Font(bold=True, size=14)
    ws.merge_cells("A1:F1")

    # 表头
    headers = ["题号", "LTC 模块", "题干", "题型", "选项 / 填空说明", "客户答案"]
    for col, h in enumerate(headers, start=1):
        c = ws.cell(row=3, column=col)
        c.value = h
        c.font = Font(bold=True)
        c.fill = PatternFill("solid", fgColor="FFE5CC")
        c.alignment = Alignment(horizontal="center", vertical="center")

    # 数据行
    type_label = {
        "single": "单选", "multi": "多选", "rating": "分级量表",
        "number": "数值", "text": "短文本", "node_pick": "节点勾选",
    }
    row_idx = 4
    global_idx = 0
    for it in items:
        global_idx += 1
        module_label = (ltc_label_lookup or {}).get(it.get("ltc_module_key", ""), it.get("ltc_module_key", ""))
        question = it.get("question", "")
        if it.get("required"):
            question += "  *"
        opt_text_lines = _option_lines(it)
        opt_text = "\n".join(opt_text_lines)

        ws.cell(row=row_idx, column=1, value=global_idx)
        ws.cell(row=row_idx, column=2, value=module_label)
        ws.cell(row=row_idx, column=3, value=question).alignment = Alignment(wrap_text=True, vertical="top")
        ws.cell(row=row_idx, column=4, value=type_label.get(it.get("type", ""), it.get("type", "")))
        ws.cell(row=row_idx, column=5, value=opt_text).alignment = Alignment(wrap_text=True, vertical="top")
        ws.cell(row=row_idx, column=6, value="").alignment = Alignment(wrap_text=True, vertical="top")
        row_idx += 1

    # 列宽
    widths = {"A": 6, "B": 18, "C": 50, "D": 12, "E": 38, "F": 30}
    for col, w in widths.items():
        ws.column_dimensions[col].width = w

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── HTML(打印用,前端用浏览器另存为 PDF)───────────────────────────────────

def export_html(
    *,
    project_name: str,
    role: str,
    items: list[dict],
    ltc_label_lookup: dict[str, str] | None = None,
) -> str:
    """生成可直接 window.print() 的 HTML。客户用浏览器「另存为 PDF」即得 PDF。"""
    groups = group_by_ltc_module(items)
    role_lbl = role_label(role)
    proj = project_name or "项目"

    parts: list[str] = []
    parts.append("<!doctype html><html lang='zh-CN'><head><meta charset='utf-8'>")
    parts.append(f"<title>{proj} · 会前调研问卷 · {role_lbl}</title>")
    parts.append("""<style>
body { font-family: -apple-system, "PingFang SC", "Microsoft YaHei", sans-serif;
       max-width: 760px; margin: 24px auto; padding: 0 24px; color: #222; line-height: 1.6; }
h1 { text-align: center; font-size: 22px; margin-bottom: 4px; }
.role { text-align: center; color: #666; font-size: 14px; margin-bottom: 24px; }
.intro { background: #fafafa; border-left: 3px solid #d96400; padding: 12px 16px; font-size: 13px; color: #555; margin-bottom: 24px; }
h2 { font-size: 16px; border-bottom: 1px solid #ddd; padding-bottom: 6px; margin-top: 28px; }
.q { margin: 16px 0 18px 0; }
.q-stem { font-weight: 600; font-size: 14px; }
.required { color: #d33; margin-left: 4px; }
.hint { color: #888; font-size: 12px; font-style: italic; margin: 4px 0 8px; }
.opts { margin: 6px 0 0 16px; font-size: 13px; }
.opts li { list-style: none; margin: 3px 0; }
.opts li::before { content: "☐ "; margin-right: 4px; }
.text-blank { font-size: 13px; color: #999; margin: 6px 0; }
.print-btn { position: fixed; top: 12px; right: 12px; background: #d96400; color: white;
             border: 0; padding: 8px 16px; border-radius: 6px; cursor: pointer; font-size: 13px; }
@media print { .print-btn { display: none; } body { margin: 0; padding: 0 16px; } }
</style></head><body>""")
    parts.append('<button class="print-btn" onclick="window.print()">打印 / 另存为 PDF</button>')
    parts.append(f"<h1>{proj} · 会前调研问卷</h1>")
    parts.append(f'<div class="role">受访角色:{role_lbl}</div>')
    parts.append(
        '<div class="intro">尊敬的同事您好,本问卷用于在调研会议前了解贵司的现状,'
        '帮助实施团队提前准备针对性的访谈话题。请您按当前真实情况勾选 / 填写;选项不全可'
        '在「其他」处补充,实在不适用的题可勾「不适用」。约 5-10 分钟可完成。</div>'
    )

    if not groups:
        parts.append('<p style="color:#999;text-align:center;padding:40px 0;">(本角色暂无会前题目)</p>')

    global_idx = 0
    for ltc_key, sub_items in groups:
        module_label = (ltc_label_lookup or {}).get(ltc_key, ltc_key)
        parts.append(f"<h2>{_html_escape(module_label)}</h2>")
        for it in sub_items:
            global_idx += 1
            stem = _html_escape(it.get("question", ""))
            req_html = '<span class="required">*</span>' if it.get("required") else ""
            parts.append('<div class="q">')
            parts.append(f'<div class="q-stem">{global_idx}. {stem}{req_html}</div>')
            if it.get("hint"):
                parts.append(f'<div class="hint">{_html_escape(it["hint"])}</div>')

            t = it.get("type")
            if t in ("single", "multi", "node_pick"):
                parts.append('<ul class="opts">')
                for o in it.get("options", []):
                    parts.append(f'<li>{_html_escape(o.get("label", ""))}</li>')
                parts.append("</ul>")
            elif t == "rating":
                scale = it.get("rating_scale", 5)
                ticks = "  ".join(f"<span style='display:inline-block;width:24px;height:24px;border:1px solid #ccc;border-radius:50%;text-align:center;line-height:24px;'>{i}</span>" for i in range(1, scale + 1))
                parts.append(f'<div class="text-blank">请圈选: {ticks}</div>')
            elif t == "number":
                unit = it.get("number_unit", "")
                suf = f" ({_html_escape(unit)})" if unit else ""
                parts.append(f'<div class="text-blank">请填数值{suf}:_______________________</div>')
            else:  # text
                parts.append('<div class="text-blank">答:____________________________________________________</div>')
                parts.append('<div class="text-blank">&nbsp;&nbsp;&nbsp;&nbsp;____________________________________________________</div>')

            parts.append("</div>")

    parts.append('<p style="text-align:center;margin-top:40px;color:#666;">感谢您的耐心填写,期待会上深入交流!</p>')
    parts.append("</body></html>")
    return "".join(parts)


def _html_escape(s: str) -> str:
    return (
        (s or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )
