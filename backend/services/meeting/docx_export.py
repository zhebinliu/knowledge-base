"""会议纪要按「02003【推荐】会议纪要模板.docx」生成 docx。

调用方:`/api/meeting/{id}/export-docx` 端点。

模板表格结构(单表 21 行 × 7 列,大量 cell merge):
  Row 0: 标题(跨行)
  Row 1: 会议名称 | 召集人员
  Row 2: 会议时间 | 会议地点
  Row 3: 会议主持 | 会议记录
  Row 4: 会议形式
  Row 5: 参会人员
  Row 6: 会议主题及内容(标题行)
  Row 7: 会议主题及内容(大块文本 — 我们填 summary + key_points)
  Row 8: 待办项(标题行)
  Row 9: 序号 | 事项 | 负责人 | 备注
  Row 10-14: 待办项 5 行(超出 5 条自动 append)
  Row 15: 待确认项(标题行)
  Row 16: 序号 | 事项 | 负责人 | 备注
  Row 17-20: 待确认项 4 行(超出自动 append)
"""
from __future__ import annotations

import copy
import io
from pathlib import Path
from typing import Any, Optional

import structlog
from docx import Document
from docx.shared import Pt
from docx.table import _Row, Table

logger = structlog.get_logger()

_TEMPLATE_PATH = Path(__file__).parent / "templates" / "minutes_template.docx"


def _set_cell_text(cell, text: str, force_size: Optional[Pt] = None) -> None:
    """安全替换单元格文字,保留 cell 第一个 paragraph 的字体格式。

    text 含 \\n 时按 paragraph 拆分,每段独立成段(模板的换行风格)。
    force_size:不为 None 时强制用该字号,适合 R7 那种模板 cell 用了大号
    标题字体但我们要写正文的场景(2026-05-12 修复用户反馈"主题内容字体太大")。
    """
    text = text or ""
    p0 = cell.paragraphs[0]
    run_fmt = None
    if p0.runs:
        run_fmt = p0.runs[0].font

    # 删 p0 之后的所有 paragraph(模板某些 cell 预留几十个空段)
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    # 清空 p0 的所有 run
    for r in list(p0.runs):
        r._element.getparent().remove(r._element)

    parts = text.split("\n") if text else [""]

    def _apply_fmt(run) -> None:
        # force_size 优先于继承
        if force_size is not None:
            run.font.size = force_size
        elif run_fmt is not None and run_fmt.size is not None:
            run.font.size = run_fmt.size
        if run_fmt is not None and run_fmt.name:
            run.font.name = run_fmt.name

    r = p0.add_run(parts[0])
    _apply_fmt(r)

    for extra in parts[1:]:
        p = cell.add_paragraph(extra)
        if p.runs:
            _apply_fmt(p.runs[0])


def _ensure_rows(table: Table, after_idx: int, needed_count: int, template_row_idx: int) -> list[_Row]:
    """确保从 after_idx + 1 开始有 needed_count 行;不够就复制 template_row_idx 行追加。

    返回这 needed_count 行的 _Row 对象列表。

    注意:python-docx 没有"在中间插入行"API,我们用底层 XML 操作:
      复制 template_row 的 <w:tr> XML,insert 到 after_idx 之后。
    """
    rows = list(table.rows)
    target_rows: list[_Row] = []
    template_row = rows[template_row_idx]

    # 当前实际可用行数(after_idx + 1 起)
    existing_after = len(rows) - (after_idx + 1)

    if existing_after >= needed_count:
        # 用现有行
        for i in range(needed_count):
            target_rows.append(rows[after_idx + 1 + i])
    else:
        # 先用现有行,再 append 新行
        for i in range(existing_after):
            target_rows.append(rows[after_idx + 1 + i])
        # 追加缺的行
        to_add = needed_count - existing_after
        for _ in range(to_add):
            new_tr = copy.deepcopy(template_row._tr)
            # 清空 cell 内容,只保留格式骨架
            for tc in new_tr.iter():
                if tc.tag.endswith("}t"):
                    tc.text = ""
            template_row._tr.addnext(new_tr)
            # template_row 的位置不变(addnext 是插到它后面),所以下次还是用同一个 template
        # 重新读 rows
        rows = list(table.rows)
        for i in range(needed_count):
            target_rows.append(rows[after_idx + 1 + i])

    return target_rows


def _format_action_items(items: list[dict]) -> list[tuple[str, str, str, str]]:
    """`待办项` 表格行 (序号, 事项, 负责人, 备注)。"""
    out = []
    for idx, a in enumerate(items, start=1):
        if not isinstance(a, dict):
            continue
        task = (a.get("task") or "").strip()
        owner = (a.get("owner") or "").strip()
        deadline = (a.get("deadline") or "").strip()
        priority = (a.get("priority") or "").strip()
        remark_bits = []
        if deadline:
            remark_bits.append(f"截止:{deadline}")
        if priority:
            remark_bits.append({"high": "高优", "medium": "中优", "low": "低优"}.get(priority, priority))
        if a.get("remark"):
            remark_bits.append(str(a["remark"]))
        out.append((str(idx), task, owner, " · ".join(remark_bits)))
    return out


def _format_unresolved(items: list[dict]) -> list[tuple[str, str, str, str]]:
    """`待确认项` 表格行 (序号, 事项, 负责人, 备注)。"""
    out = []
    for idx, u in enumerate(items, start=1):
        if not isinstance(u, dict):
            continue
        issue = (u.get("issue") or "").strip()
        owner = (u.get("owner") or "").strip()
        reason = (u.get("reason") or "").strip()
        remark_bits = []
        if reason:
            remark_bits.append(f"原因:{reason}")
        if u.get("remark"):
            remark_bits.append(str(u["remark"]))
        out.append((str(idx), issue, owner, " · ".join(remark_bits)))
    return out


def _format_main_content(summary: str, key_points: list[dict], decisions: list[dict]) -> str:
    """会议主题及内容主体(模板 row 7)。"""
    lines: list[str] = []
    if summary:
        lines.append("【会议摘要】")
        lines.append(summary)
        lines.append("")
    if key_points:
        lines.append("【会议主题及内容】")
        for i, kp in enumerate(key_points, start=1):
            if not isinstance(kp, dict):
                continue
            topic = (kp.get("topic") or "").strip()
            content = (kp.get("content") or "").strip()
            lines.append(f"{i}) {topic}")
            if content:
                lines.append(content)
            lines.append("")
    if decisions:
        lines.append("【决议事项】")
        for d in decisions:
            if not isinstance(d, dict):
                continue
            c = (d.get("content") or "").strip()
            o = (d.get("owner") or "").strip()
            line = f"• {c}"
            if o:
                line += f"(负责人:{o})"
            lines.append(line)
        lines.append("")
    return "\n".join(lines).rstrip()


def render_minutes_docx(
    meeting_title: str,
    minutes: dict[str, Any],
    fallback_time: str = "",
    fallback_attendees: str = "",
) -> bytes:
    """根据 minutes 数据 + 模板生成 docx,返回 bytes。"""
    if not _TEMPLATE_PATH.exists():
        raise RuntimeError(f"会议纪要模板缺失:{_TEMPLATE_PATH}")

    doc = Document(str(_TEMPLATE_PATH))
    if not doc.tables:
        raise RuntimeError("模板格式异常:未找到表格")

    table = doc.tables[0]

    # 取 minutes 字段(都带兜底)
    m_title = (minutes.get("meeting_title") or meeting_title or "").strip()
    m_time = (minutes.get("meeting_time") or fallback_time or "").strip()
    m_location = (minutes.get("meeting_location") or "").strip()
    m_host = (minutes.get("meeting_host") or "").strip()
    m_recorder = (minutes.get("meeting_recorder") or "").strip()
    m_format = (minutes.get("meeting_format") or "").strip()
    m_organizer = (minutes.get("organizer") or "").strip()
    attendees = minutes.get("attendees") or []
    if isinstance(attendees, list):
        attendees_text = "\n".join(str(a) for a in attendees if a)
    else:
        attendees_text = str(attendees)
    if not attendees_text and fallback_attendees:
        attendees_text = fallback_attendees

    rows = list(table.rows)

    # Row 0: 标题
    _set_cell_text(rows[0].cells[0], m_title or "会议纪要")

    # Row 1: 会议名称 | 召集人员
    _set_cell_text(rows[1].cells[2], m_title)
    _set_cell_text(rows[1].cells[5], m_organizer)

    # Row 2: 会议时间 | 会议地点
    _set_cell_text(rows[2].cells[2], m_time)
    _set_cell_text(rows[2].cells[5], m_location)

    # Row 3: 主持 | 记录
    _set_cell_text(rows[3].cells[2], m_host)
    _set_cell_text(rows[3].cells[5], m_recorder)

    # Row 4: 会议形式(可能跨列)
    _set_cell_text(rows[4].cells[2], m_format)

    # Row 5: 参会人员
    _set_cell_text(rows[5].cells[2], attendees_text)

    # Row 7: 会议主题及内容(大块)— 强制用 10.5pt 五号,模板原 cell 字体偏大
    main_content = _format_main_content(
        summary=(minutes.get("summary") or "").strip(),
        key_points=minutes.get("key_points") or [],
        decisions=minutes.get("decisions") or [],
    )
    _set_cell_text(rows[7].cells[0], main_content, force_size=Pt(10.5))

    # 待办项:Row 10-14 是 5 行默认,超出 append。表头在 row 9。
    action_rows_data = _format_action_items(minutes.get("action_items") or [])
    # 至少留 1 行(空也保留模板原行)
    needed_action = max(len(action_rows_data), 1)
    action_target_rows = _ensure_rows(table, after_idx=9, needed_count=needed_action, template_row_idx=10)
    for trow, data in zip(action_target_rows, action_rows_data):
        _set_cell_text(trow.cells[0], data[0])
        _set_cell_text(trow.cells[1], data[1])
        _set_cell_text(trow.cells[3], data[2])
        _set_cell_text(trow.cells[6], data[3])
    # 没数据时把序号写"-"提示
    for trow in action_target_rows[len(action_rows_data):]:
        _set_cell_text(trow.cells[0], "-")
        _set_cell_text(trow.cells[1], "")
        _set_cell_text(trow.cells[3], "")
        _set_cell_text(trow.cells[6], "")

    # 待确认项:_ensure_rows 已经在前面 append 过,需要重新读 rows 找标题位置
    rows = list(table.rows)
    # 找"待确认项"那一行(text 包含)
    confirm_header_idx = None
    for i, r in enumerate(rows):
        if "待确认项" in r.cells[0].text:
            confirm_header_idx = i
            break

    if confirm_header_idx is not None:
        # 表头行 = confirm_header_idx + 1,数据起始 = confirm_header_idx + 2
        unresolved_data = _format_unresolved(minutes.get("unresolved") or [])
        needed_unres = max(len(unresolved_data), 1)
        unres_template_idx = confirm_header_idx + 2
        # 确保 template_row 存在
        if unres_template_idx < len(rows):
            unres_target_rows = _ensure_rows(
                table,
                after_idx=confirm_header_idx + 1,
                needed_count=needed_unres,
                template_row_idx=unres_template_idx,
            )
            for trow, data in zip(unres_target_rows, unresolved_data):
                _set_cell_text(trow.cells[0], data[0])
                _set_cell_text(trow.cells[1], data[1])
                _set_cell_text(trow.cells[3], data[2])
                _set_cell_text(trow.cells[6], data[3])
            for trow in unres_target_rows[len(unresolved_data):]:
                _set_cell_text(trow.cells[0], "-")
                _set_cell_text(trow.cells[1], "")
                _set_cell_text(trow.cells[3], "")
                _set_cell_text(trow.cells[6], "")

    # 保存到 bytes
    buf = io.BytesIO()
    doc.save(buf)
    out = buf.getvalue()
    logger.info("minutes_docx_rendered", title=m_title[:40], bytes=len(out))
    return out
